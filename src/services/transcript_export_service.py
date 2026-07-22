"""Exportação da transcrição (texto editável) para TXT/DOCX/MD e para
prompts prontos de IA (ChatGPT/Claude).

Responsabilidade única: formatar e gravar texto em disco. Não sabe nada de
GUI nem de como a transcrição foi gerada (Whisper) — recebe sempre o
`Transcription` já pronto.
"""
from __future__ import annotations

from pathlib import Path

from src.models.domain import Transcription
from src.utils.logger import get_logger
from src.utils.paths import sanitize_filename

logger = get_logger("transcript_export")

# Pausas mais longas que isso viram quebra de parágrafo no texto legível
# (ver `README` da funcionalidade: facilita a leitura de um roteiro corrido).
PARAGRAPH_GAP_SECONDS = 1.0

SCRIPT_PROMPT_HEADER = """Você é um roteirista profissional para YouTube.

Reescreva este roteiro completamente.

Regras:
- Não copie frases.
- Mantenha a ordem dos acontecimentos.
- Mantenha a duração aproximada.
- Adicione suspense e transições naturais.
- Escreva em português brasileiro.
- Produza um roteiro pronto para narração por IA.

Transcrição:

"""


class TranscriptExportService:
    """Formata e exporta transcrições em texto legível."""

    # ------------------------------------------------------------------ #
    # Formatação
    # ------------------------------------------------------------------ #
    @staticmethod
    def format_readable(
        transcription: Transcription, pause_threshold: float = PARAGRAPH_GAP_SECONDS,
    ) -> str:
        """Texto corrido da narração, preservando a ordem original e
        quebrando parágrafo sempre que houver uma pausa maior que
        `pause_threshold` segundos entre dois trechos."""
        paragraphs: list[str] = []
        current: list[str] = []
        prev_end: float | None = None
        for seg in transcription.segments:
            text = seg.text.strip()
            if not text:
                continue
            if prev_end is not None and (seg.start - prev_end) > pause_threshold and current:
                paragraphs.append(" ".join(current))
                current = []
            current.append(text)
            prev_end = seg.end
        if current:
            paragraphs.append(" ".join(current))
        return "\n\n".join(paragraphs)

    @staticmethod
    def build_script_prompt(text: str) -> str:
        """Monta o arquivo pronto pra colar no ChatGPT/Claude: instruções de
        roteirista + a transcrição (limpa, sem timestamps/IDs/metadados)."""
        return f"{SCRIPT_PROMPT_HEADER}{text.strip()}\n"

    # ------------------------------------------------------------------ #
    # Nome de arquivo
    # ------------------------------------------------------------------ #
    @staticmethod
    def default_filename(video_title: str, suffix: str, extension: str) -> str:
        stem = sanitize_filename(video_title)
        return f"{stem}_{suffix}.{extension}"

    # ------------------------------------------------------------------ #
    # Gravação em disco
    # ------------------------------------------------------------------ #
    @staticmethod
    def save_txt(text: str, path: str | Path) -> Path:
        path = Path(path)
        path.write_text(text, encoding="utf-8")
        logger.info("Transcrição exportada em TXT: %s", path)
        return path

    @staticmethod
    def save_markdown(text: str, path: str | Path, title: str) -> Path:
        path = Path(path)
        body = f"# Transcrição — {title}\n\n{text}\n"
        path.write_text(body, encoding="utf-8")
        logger.info("Transcrição exportada em Markdown: %s", path)
        return path

    @staticmethod
    def save_docx(text: str, path: str | Path, title: str) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx não instalado. Rode: pip install python-docx"
            ) from exc
        path = Path(path)
        document = Document()
        document.add_heading(title, level=1)
        for paragraph in text.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                document.add_paragraph(paragraph)
        document.save(str(path))
        logger.info("Transcrição exportada em DOCX: %s", path)
        return path
